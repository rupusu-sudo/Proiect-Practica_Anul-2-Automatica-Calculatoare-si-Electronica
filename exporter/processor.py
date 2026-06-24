import os
import uuid
import logging
from typing import Dict, Any
from datetime import datetime

# Importuri reportlab și python-docx
from reportlab.lib.pagesizes import letter  # pyrefly: ignore [missing-import]
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # pyrefly: ignore [missing-import]
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # pyrefly: ignore [missing-import]
from docx import Document  # pyrefly: ignore [missing-import]

logger = logging.getLogger("a2a.exporter")

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


class ExportProcessor:
    """Generează fișiere PDF și DOCX concomitent pentru rapoartele academice."""

    def process_task(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        title = payload.get("title", "Raport de Cercetare")
        report_text = payload.get("report", "")
        summary = payload.get("summary", "")
        key_points = payload.get("key_points", [])
        sources = payload.get("sources", [])
        trust_score = payload.get("trust_score", 70)
        timestamp = payload.get("timestamp", datetime.now().isoformat())
        workflow_metadata = payload.get("workflow_metadata", {})

        # Generare nume fișiere unice
        unique_id = uuid.uuid4().hex[:12]
        pdf_filename = f"cercetare_{unique_id}.pdf"
        docx_filename = f"cercetare_{unique_id}.docx"
        
        pdf_filepath = os.path.join(EXPORT_DIR, pdf_filename)
        docx_filepath = os.path.join(EXPORT_DIR, docx_filename)

        # Generare ambele formate
        self._generate_pdf(pdf_filepath, title, report_text, summary, key_points, sources, trust_score, timestamp, workflow_metadata)
        self._generate_docx(docx_filepath, title, report_text, summary, key_points, sources, trust_score, timestamp, workflow_metadata)

        logger.info(f"Documente exportate cu succes la: {pdf_filename} și {docx_filename}")

        return {
            "pdf_filename": pdf_filename,
            "docx_filename": docx_filename,
            "pdf_filepath": pdf_filepath,
            "docx_filepath": docx_filepath,
            "trust_score": trust_score,
            "verified_sources": payload.get("verified_sources", sources),
            "rejected_sources": payload.get("rejected_sources", [])
        }

    def _generate_pdf(self, filepath: str, title: str, report_text: str, summary: str, key_points: list, sources: list, trust_score: int, timestamp: str, metadata: dict):
        """Generează un document PDF academic formatat."""
        doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        # Creare stiluri personalizate
        title_style = ParagraphStyle(
            'AcademicTitle',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            spaceAfter=12,
            textColor='#2563eb'
        )
        heading_style = ParagraphStyle(
            'AcademicHeading2',
            parent=styles['Heading2'],
            fontSize=13,
            leading=16,
            spaceBefore=14,
            spaceAfter=6,
            textColor='#1f2937'
        )
        body_style = ParagraphStyle(
            'AcademicBody',
            parent=styles['BodyText'],
            fontSize=10,
            leading=14,
            spaceAfter=8,
            textColor='#374151'
        )
        meta_style = ParagraphStyle(
            'AcademicMeta',
            parent=styles['Normal'],
            fontSize=8,
            leading=11,
            textColor='#6b7280'
        )

        story = []
        
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 10))

        # Metadate
        story.append(Paragraph(f"<b>Data Generare:</b> {timestamp}", meta_style))
        story.append(Paragraph(f"<b>Scor Credibilitate Surse:</b> {trust_score} / 100", meta_style))
        if metadata:
            meta_str = " | ".join([f"{k}: {v}" for k, v in metadata.items()])
            story.append(Paragraph(f"<b>Metadata A2A Workflow:</b> {meta_str}", meta_style))
        story.append(Spacer(1, 15))

        # Rezumat
        if summary:
            story.append(Paragraph("Rezumat", heading_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 10))

        # Puncte Cheie
        if key_points:
            story.append(Paragraph("Puncte Cheie", heading_style))
            for pt in key_points:
                story.append(Paragraph(f"• {pt}", body_style))
            story.append(Spacer(1, 10))

        # Raport Structurat (cu paragrafe)
        if report_text:
            story.append(Paragraph("Raport Academic Structurat", heading_style))
            for paragraph in report_text.split("\n\n"):
                if paragraph.strip():
                    story.append(Paragraph(paragraph.replace("\n", "<br/>"), body_style))

        # Referințe / Surse
        if sources:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Surse Consultate", heading_style))
            for src in sources:
                story.append(Paragraph(f"• {src}", body_style))

        doc.build(story)

    def _generate_docx(self, filepath: str, title: str, report_text: str, summary: str, key_points: list, sources: list, trust_score: int, timestamp: str, metadata: dict):
        """Generează un document Word .docx formatat."""
        doc = Document()
        
        doc.add_heading(title, level=1)
        
        p_meta = doc.add_paragraph()
        p_meta.add_run(f"Data Generare: {timestamp}\n").italic = True
        p_meta.add_run(f"Scor Credibilitate Surse: {trust_score} / 100\n").bold = True
        if metadata:
            meta_str = " | ".join([f"{k}: {v}" for k, v in metadata.items()])
            p_meta.add_run(f"Metadata A2A Workflow: {meta_str}").italic = True
            
        doc.add_paragraph()

        # Rezumat
        if summary:
            doc.add_heading("Rezumat", level=2)
            doc.add_paragraph(summary)

        # Puncte Cheie
        if key_points:
            doc.add_heading("Puncte Cheie", level=2)
            for pt in key_points:
                doc.add_paragraph(pt, style='List Bullet')

        # Raport Structurat
        if report_text:
            doc.add_heading("Raport Academic Structurat", level=2)
            for paragraph in report_text.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)

        # Referințe / Surse
        if sources:
            doc.add_heading("Surse Consultate", level=2)
            for src in sources:
                doc.add_paragraph(src, style='List Bullet')

        doc.save(filepath)
